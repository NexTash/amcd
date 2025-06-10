# # Copyright (c) 2025, Umer and contributors
# # For license information, please see license.txt

import os
from docx import Document as DocxDocument
import frappe
from frappe.model.document import Document
from frappe.utils.file_manager import save_file


class CLOCheck(Document):
    def validate(self):
        def extract_questions_and_bold_words_from_docx(file_path):
            doc = DocxDocument(file_path)
            questions = []
            bold_words = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    questions.append(paragraph.text.strip())

                    # Extract bold words from this paragraph
                    for run in paragraph.runs:
                        if run.bold and run.text.strip():
                            bold_words.append(run.text.strip())

            return questions, bold_words

        def process_file(file_url):
            full_path = frappe.utils.file_manager.get_file_path(file_url)
            _, ext = os.path.splitext(full_path)
            if ext.lower() == ".docx":
                return extract_questions_and_bold_words_from_docx(full_path)
            else:
                frappe.throw("Unsupported file type. Please upload a .docx file.")

        def append_to_child_table(parent_doc, questions, bold_words):
            parent_doc.set("paper_questions", [])
            for i, question in enumerate(questions):
                clo_word = bold_words[i] if i < len(bold_words) else ""
                parent_doc.append("paper_questions", {
                    "question": question,
                    "clo_word": clo_word
                })

        if self.question_paper:
            questions, bold_words = process_file(self.question_paper)
            append_to_child_table(self, questions, bold_words)
        
        def get_clo_parent_name(clo_word):
            # Use Frappe's database query to fetch the parent name
            clo_doc = frappe.db.sql("""
                SELECT parent 
                FROM `tabCLO Words` 
                WHERE word = %s
            """, (clo_word,), as_dict=True)
        
            return clo_doc[0].parent if clo_doc else ""
        
        for question_row in self.paper_questions:
            clo = get_clo_parent_name(question_row.clo_word)
            question_row.clo = clo
    
    @frappe.whitelist()
    def download_clo_file(self):
        questions = self.paper_questions or []
        
        doc = DocxDocument()

        for row in questions:
            doc.add_paragraph(row.question or "")
            
            if row.clo:
                doc.add_paragraph(row.clo or "")
            else:
                doc.add_paragraph("NO CLO Found")

        file_path = "/tmp/CLO_File.docx"
        doc.save(file_path)

        file_name = "CLO_File.docx"
        file_data = open(file_path, "rb").read()
        new_file = save_file(file_name, file_data, self.doctype, self.name)

        os.remove(file_path)

        return new_file